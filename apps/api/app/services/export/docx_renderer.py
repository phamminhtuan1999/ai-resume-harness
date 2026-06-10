"""ATS-safe DOCX renderer for the Draft CV (US-042, reworked by US-044).

Renders the same shared render model as the PDF (US-041) into a single-column
Word document using real heading paragraphs and plain list bullets — no tables,
text boxes, images, or header/footer content — so it parses cleanly in ATS and
opens without repair warnings. python-docx is pure Python (no native deps) and
preserves full Unicode (names with diacritics survive intact).

US-044 (decision 0014 §2): the ``font_profile`` maps to exactly **one**
universally-available font name applied to body and heading styles — OOXML has
no per-run fallback chain and python-docx cannot embed fonts, so naming a font
every machine has is the only predictable option.
"""

from __future__ import annotations

import io
from typing import Any

from app.services.export.fonts import resolve_font_profile
from app.services.export.options import RenderOptions
from app.services.export.render_config import get_render_config


def render_docx(
    render_model: dict[str, Any], options: RenderOptions | None = None
) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise RuntimeError("DOCX renderer dependency is unavailable.") from exc

    opts = options or RenderOptions()
    profile = resolve_font_profile(opts.font_profile)
    config = get_render_config(opts.page_target, opts.density)
    doc = Document()
    _apply_typography(doc, Pt, profile.docx_font, config.body_pt)

    contact = render_model.get("contact") or {}
    name = (contact.get("full_name") or "").strip()
    if name:
        heading = doc.add_heading(name, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact_line = "  |  ".join(
        str(b)
        for b in (
            contact.get("email"),
            contact.get("phone"),
            contact.get("location"),
            contact.get("linkedin_url"),
            contact.get("github_url"),
            contact.get("portfolio_url"),
        )
        if b
    )
    if contact_line:
        doc.add_paragraph(contact_line)

    summary = (render_model.get("professional_summary") or "").strip()
    if summary:
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(summary)

    skills = render_model.get("skills") or []
    if skills:
        doc.add_heading("Skills", level=1)
        for group in skills:
            items = ", ".join(group.get("items") or [])
            if not items:
                continue
            para = doc.add_paragraph()
            run = para.add_run(f"{group.get('category') or 'Skills'}: ")
            run.bold = True
            para.add_run(items)

    _experience_section(doc, "Work Experience", render_model.get("work_experience") or [])
    _projects_section(doc, render_model.get("projects") or [])
    _education_section(doc, render_model.get("education") or [])
    _certifications_section(doc, render_model.get("certifications") or [])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _apply_typography(doc: Any, pt: Any, font_name: str, body_pt: float) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = pt(body_pt)
    # Headings inherit theme fonts unless overridden per style; the profile
    # font must win everywhere for a consistent, ATS-predictable document.
    for style_name in ("Title", "Heading 1"):
        try:
            doc.styles[style_name].font.name = font_name
        except KeyError:  # pragma: no cover - styles exist in default template
            continue


def _entry_heading(doc: Any, primary: str, secondary: str, dates: str) -> None:
    para = doc.add_paragraph()
    title = primary if not secondary else f"{primary} — {secondary}"
    run = para.add_run(title)
    run.bold = True
    if dates:
        para.add_run(f"  ({dates})")


def _bullets(doc: Any, bullets: list[str]) -> None:
    for text in bullets:
        if text.strip():
            doc.add_paragraph(text, style="List Bullet")


def _experience_section(doc: Any, heading: str, entries: list[dict[str, Any]]) -> None:
    visible = [e for e in entries if (e.get("company") or e.get("title") or e.get("bullets"))]
    if not visible:
        return
    doc.add_heading(heading, level=1)
    for entry in visible:
        _entry_heading(
            doc,
            entry.get("title") or "",
            entry.get("company") or "",
            _date_range(entry.get("start_date"), entry.get("end_date")),
        )
        if entry.get("location"):
            doc.add_paragraph(str(entry["location"]))
        _bullets(doc, entry.get("bullets") or [])


def _projects_section(doc: Any, entries: list[dict[str, Any]]) -> None:
    visible = [e for e in entries if (e.get("name") or e.get("bullets"))]
    if not visible:
        return
    doc.add_heading("Projects", level=1)
    for entry in visible:
        _entry_heading(doc, entry.get("name") or "", ", ".join(entry.get("tech_stack") or []), "")
        if entry.get("description"):
            doc.add_paragraph(str(entry["description"]))
        _bullets(doc, entry.get("bullets") or [])


def _education_section(doc: Any, entries: list[dict[str, Any]]) -> None:
    visible = [e for e in entries if e.get("school")]
    if not visible:
        return
    doc.add_heading("Education", level=1)
    for entry in visible:
        line = ", ".join(b for b in (entry.get("degree"), entry.get("field")) if b)
        _entry_heading(
            doc,
            entry.get("school") or "",
            line,
            _date_range(entry.get("start_date"), entry.get("end_date")),
        )
        if entry.get("details"):
            doc.add_paragraph(str(entry["details"]))


def _certifications_section(doc: Any, entries: list[dict[str, Any]]) -> None:
    visible = [e for e in entries if e.get("name")]
    if not visible:
        return
    doc.add_heading("Certifications", level=1)
    for entry in visible:
        line = " — ".join(
            str(b) for b in (entry.get("name"), entry.get("issuer"), entry.get("date")) if b
        )
        doc.add_paragraph(line, style="List Bullet")


def _date_range(start: Any, end: Any) -> str:
    start = str(start).strip() if start else ""
    end = str(end).strip() if end else ""
    if start and end:
        return f"{start} – {end}"
    return start or end or ""
