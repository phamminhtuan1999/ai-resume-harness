"""US-041/US-042 export tests: render-model gating, PDF + DOCX content, parity.

The render model is the single gating boundary, so the highest-value proof is
content assertion on the produced files: approved/safe text must appear, and
do_not_use_yet / pending / rejected text must never appear, in either format.
"""

from __future__ import annotations

import io
import re

from app.services.export.docx_renderer import render_docx
from app.services.export.pdf_renderer import render_pdf
from app.services.export.render_model import (
    build_render_model,
    filename_slug,
    is_empty_cv,
    is_renderable,
    pending_review_count,
    renderable_bullet_count,
)


def _bullet(text: str, status: str, action: str = "pending") -> dict:
    return {
        "id": f"b-{text.split()[1]}",
        "text": text,
        "source_evidence": "ev",
        "truth_guard_status": status,
        "keywords_used": [],
        "user_action": action,
    }


def _cv_json() -> dict:
    """One bullet of every truth-guard x user_action combination."""
    return {
        "candidate": {"full_name": "Dana Engineer", "email": "dana@example.com"},
        "target_job": {"company": "Acme AI", "title": "Senior AI Engineer"},
        "professional_summary": "Senior engineer summary line.",
        "skills": [
            {"category": "Backend", "items": ["FastAPI"]},
            {"category": "Empty", "items": []},  # dropped by the render model
        ],
        "work_experience": [
            {
                "company": "Acme",
                "title": "Engineer",
                "location": "Remote",
                "start_date": "2020",
                "end_date": "2024",
                "bullets": [
                    _bullet("Engineered Approvedalpha services.", "safe_to_use"),
                    _bullet("Designed Approvedbeta schemas.", "needs_confirmation", "approved"),
                    _bullet("Built Pendinggamma pipelines.", "needs_confirmation", "pending"),
                    _bullet("Created Rejecteddelta tools.", "needs_confirmation", "rejected"),
                    _bullet("Improved Forbiddenepsilon throughput.", "do_not_use_yet"),
                ],
            }
        ],
        "projects": [],
        "education": [{"school": "State University", "degree": "BSc", "field": "CS"}],
        "certifications": [{"name": "AWS SAA", "issuer": "Amazon", "date": "2022"}],
    }


_RENDERABLE = ("Approvedalpha", "Approvedbeta")
_EXCLUDED = ("Pendinggamma", "Rejecteddelta", "Forbiddenepsilon")


def _compact(text: str) -> str:
    return re.sub(r"\s+", "", text)


# --- render model gating --------------------------------------------------------


def test_is_renderable_predicate() -> None:
    assert is_renderable(_bullet("x y", "safe_to_use"))
    assert is_renderable(_bullet("x y", "needs_confirmation", "approved"))
    assert not is_renderable(_bullet("x y", "needs_confirmation", "pending"))
    assert not is_renderable(_bullet("x y", "needs_confirmation", "rejected"))
    assert not is_renderable(_bullet("x y", "do_not_use_yet"))


def test_render_model_filters_and_drops_empty_groups() -> None:
    rm = build_render_model(_cv_json())
    bullets = rm["work_experience"][0]["bullets"]
    assert any("Approvedalpha" in b for b in bullets)
    assert any("Approvedbeta" in b for b in bullets)
    for excluded in _EXCLUDED:
        assert all(excluded not in b for b in bullets)
    assert renderable_bullet_count(rm) == 2
    assert [g["category"] for g in rm["skills"]] == ["Backend"]  # empty group dropped


def test_pending_review_count() -> None:
    assert pending_review_count(_cv_json()) == 1  # only Pendinggamma


def test_is_empty_cv() -> None:
    empty = {
        "professional_summary": "",
        "work_experience": [
            {"bullets": [_bullet("Improved Forbidden throughput.", "do_not_use_yet")]}
        ],
        "projects": [],
    }
    assert is_empty_cv(build_render_model(empty))
    assert not is_empty_cv(build_render_model(_cv_json()))


def test_filename_slug_is_ascii() -> None:
    rm = build_render_model({"candidate": {"full_name": "Đặng Quốc Tuấn"}})
    slug = filename_slug(rm, {"company": "Acme AI", "title": "Senior Engineer"})
    assert re.fullmatch(r"[a-z0-9-]+", slug)
    assert "/" not in slug and " " not in slug


# --- PDF ------------------------------------------------------------------------


def _pdf_text(pdf_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def test_pdf_is_valid_and_gates_content() -> None:
    pdf = render_pdf(build_render_model(_cv_json()))
    assert pdf[:5] == b"%PDF-"
    compact = _compact(_pdf_text(pdf))
    for token in _RENDERABLE:
        assert token in compact
    for token in _EXCLUDED:
        assert token not in compact


def test_pdf_handles_unicode_name_without_crashing() -> None:
    cv = _cv_json()
    cv["candidate"]["full_name"] = "Đặng Quốc Tuấn"
    pdf = render_pdf(build_render_model(cv))
    assert pdf[:5] == b"%PDF-"  # no UnicodeEncodeError from core fonts


def test_pdf_renders_real_typographic_glyphs_never_question_marks() -> None:
    """The renderer's own separators (• bullets, – date ranges, — title joins)
    and data-borne curly quotes once degraded to "?" under the latin-1 core
    fonts (intake #41). US-044's default embedded Unicode fonts render the real
    glyphs; the regression guard is that "?" never appears. (ASCII
    transliteration on the missing-asset fallback is covered in
    test_draft_cv_fonts.)"""
    cv = _cv_json()
    cv["professional_summary"] = "Shipped “zero-downtime” migrations — 99.9% uptime."
    cv["work_experience"][0]["start_date"] = "Oct 2022"
    cv["work_experience"][0]["end_date"] = "Present"
    text = _pdf_text(render_pdf(build_render_model(cv)))
    assert "?" not in text
    compact = _compact(text)
    assert "shipped“zero-downtime”migrations—99.9%uptime." in compact.lower()
    assert "(oct2022–present)" in compact.lower()  # real en dash preserved
    assert "•engineeredapprovedalpha" in compact.lower()  # real bullet preserved


# --- DOCX -----------------------------------------------------------------------


def _docx_doc(docx_bytes: bytes):
    from docx import Document

    return Document(io.BytesIO(docx_bytes))


def test_docx_is_valid_and_gates_content() -> None:
    docx = render_docx(build_render_model(_cv_json()))
    # PK zip magic (DOCX is a zip container).
    assert docx[:2] == b"PK"
    doc = _docx_doc(docx)
    text = "\n".join(p.text for p in doc.paragraphs)
    compact = _compact(text)
    for token in _RENDERABLE:
        assert token in compact
    for token in _EXCLUDED:
        assert token not in compact


def test_docx_is_ats_safe_structure() -> None:
    doc = _docx_doc(render_docx(build_render_model(_cv_json())))
    assert doc.tables == []  # no tables
    assert len(doc.inline_shapes) == 0  # no images
    # Section titles use real heading styles.
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Work Experience" in headings


def test_docx_preserves_unicode_name() -> None:
    cv = _cv_json()
    cv["candidate"]["full_name"] = "Đặng Quốc Tuấn"
    doc = _docx_doc(render_docx(build_render_model(cv)))
    assert any("Đặng Quốc Tuấn" in p.text for p in doc.paragraphs)


# --- parity ---------------------------------------------------------------------


def test_pdf_and_docx_render_the_same_gated_content() -> None:
    rm = build_render_model(_cv_json())
    pdf_compact = _compact(_pdf_text(render_pdf(rm)))
    doc = _docx_doc(render_docx(rm))
    docx_compact = _compact("\n".join(p.text for p in doc.paragraphs))
    for token in _RENDERABLE:
        assert token in pdf_compact and token in docx_compact
    for token in _EXCLUDED:
        assert token not in pdf_compact and token not in docx_compact
